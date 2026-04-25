from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"C:\Users\chen yin\Desktop\lunwen\正文\毕业论文正式稿_模板整理版_更新版.docx")
BACKUP_PATH = DOC_PATH.with_name(DOC_PATH.stem + "_加高度分层前备份.docx")


SECTION_BLOCKS = {
    "3.4.4 不同高度任务分层规划策略": [
        "3.4.4 不同高度任务分层规划策略",
        "在多无人机编队飞行任务中，若所有飞行器始终处于同一标称高度层，则在编队重构、通道穿越和任务切换阶段更容易出现局部拥挤、横向展开不足和轨迹交叉风险。为提高任务组织的灵活性，可在现有编队航路生成机制基础上引入不同高度任务分层规划策略，即在保持总体任务目标一致的前提下，为部分飞行器或子编队分配临时高度层，使其在重组阶段先通过高度分离获得额外缓冲，再逐步回归统一任务高度。",
        "该策略的核心思想是将原本完全依赖平面横向偏置的编队重构过程扩展为“平面偏置与高度分层协同”的组织过程。具体而言，可将任务执行划分为起飞集结阶段、编队稳定阶段、分层过渡阶段和任务恢复阶段。在起飞集结阶段，机群先在统一参考高度附近完成速度和航向收敛；在编队稳定阶段，维持当前任务所需的基本队形；当进入狭窄通道、障碍密集区或中途改编队场景时，根据编队角色和冲突风险将飞行器分配到不同临时高度层；在任务恢复阶段，再将各高度层逐步收敛至统一巡航高度，从而完成从“分层缓冲”到“统一执行”的转换。",
        "从任务规划实现角度看，不同高度分层规划并不要求完全重写现有航路生成框架，而是可以建立在现有任务点、重组点和编队参数生成逻辑之上。系统可首先根据任务锚点和编队类型确定参考高度，再结合左右翼位置、前后顺序或任务优先级，为各无人机施加有限幅度的临时高度偏置。随后，在任务段之间插入爬升、保持和回收三个子过程，使飞行器在进入关键区域前完成层间分离，在离开关键区域后恢复统一高度。对于固定翼平台而言，该策略尤其适用于横队和 V 字队形等横向展开较大的编队形式，因为此类队形在同一平面内更依赖额外的转弯与重组空间。",
        "需要指出的是，高度分层策略在本文中主要被视为面向多无人机任务组织的扩展规划手段，其目标是降低重构过程中的局部冲突概率、改善编队切换的可执行性，并为后续更高保真的三维协同规划研究提供接口。现阶段系统中的障碍建模与局部避障仍以二维平面关系为主，因此该策略更适合表述为“在既有任务规划框架上的分层增强机制”，而非已经完全成熟的三维全局避障系统。",
    ],
    "4.5.6 不同高度任务场景": [
        "4.5.6 不同高度任务场景",
        "为验证不同高度任务分层规划策略的可行性，可在现有动态改编队场景和障碍密集场景基础上构造高度分层任务场景。该场景保持总体任务目标点不变，但在编队切换或关键重组段引入临时高度层配置，使部分无人机在进入关键区域前先爬升至较高任务层，另一部分无人机保持原巡航高度或进入较低缓冲层，由此形成“同一水平任务目标、不同临时高度层协同执行”的组织方式。",
        "在实验设计中，可将高度分层参数设置为有限个离散等级，例如以任务参考高度为基准，分别配置上层、基准层和下层三类临时高度层，并控制层间高度差保持在固定范围内。实验时重点观察机群在进入重组区、完成分层飞行以及重新汇合时的稳定性变化，从而分析高度层策略对固定翼编队重构、局部冲突缓解和任务连续性的影响。该场景并非替代现有四类典型场景，而是作为增强型扩展场景，用于考察现有系统在由二维任务组织向准三维任务组织过渡时的适应能力。",
    ],
    "4.7.9 高度分层规划相关指标": [
        "4.7.9 高度分层规划相关指标",
        "在引入不同高度任务分层规划后，原有任务完成率、平均航迹长度、最小安全距离和编队误差等指标仍然有效，但还需补充若干能够反映分层策略效果的观察量。第一，可记录临时高度层切换次数，用以反映系统在任务执行中触发分层调度的频繁程度；第二，可统计高度恢复时间，即飞行器从进入临时高度层到重新回归统一任务高度所经历的时间；第三，可比较分层前后额外航程与估算能耗增量，用以分析高度分离带来的任务代价变化；第四，可观察层间最小垂向间隔是否满足预设安全要求，以判断高度层设置是否真正发挥空间缓冲作用。",
        "上述指标中，部分指标可直接由任务点序列与飞行状态记录计算得到，部分指标则需结合运行态截图和重复实验记录进行综合判断。通过将这些指标纳入统一评价体系，可使高度分层规划不再仅作为概念性扩展，而能够在实验设计层面获得更明确的分析口径。",
    ],
    "5.4.5 不同高度分层规划的扩展分析": [
        "5.4.5 不同高度分层规划的扩展分析",
        "结合现有系统实现基础可以看出，在多无人机任务组织框架中引入不同高度任务分层规划具有较好的扩展可行性。一方面，当前任务点和航迹点本身已包含高度信息，平均航迹长度与简化能耗模型也已将高度变化纳入估算过程；另一方面，系统已具备集结点插入、任务刷新、编队重构和阶段性任务组织能力，因此在关键任务段施加临时高度偏置，并不会与现有任务生成框架发生根本冲突。",
        "从编队执行机理来看，高度分层规划最直接的作用在于为编队切换和局部重组提供额外的空间缓冲。对于横队和 V 字队形这类横向展开明显的编队形式，若完全依赖同一平面内的横向拉开，固定翼平台往往需要更长的过渡距离和更大的转弯余度；而在引入临时高度层后，部分飞行器可以先通过垂向分离减小同平面拥挤，再在重组完成后回归统一任务高度。这样既有助于缓解局部交叉风险，也能降低因单机迟滞导致整个编队释放时机被过度推迟的概率。",
        "不过需要保持清醒认识的是，现阶段该分层策略更适合作为现有系统的增强型任务组织机制，而非完全成熟的三维避障结论。当前障碍建模和局部避障逻辑仍主要围绕二维投影关系展开，因此不同高度层更多承担“重组缓冲”和“任务分流”的作用，而不是独立完成复杂三维障碍规避。基于这一定位，在后续实验中可优先通过典型截图、任务执行过程对比和运行态观察来验证其工程价值，再视数据积累情况逐步补充定量统计分析。",
    ],
}


def find_paragraph(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"未找到段落: {text}")


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def insert_after(paragraph, text: str, style):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    wrapped = Paragraph(new_p, paragraph._parent)
    wrapped.style = style
    wrapped.add_run(text)
    return wrapped


def cleanup_existing_sections(doc: Document):
    texts_to_remove = {line for block in SECTION_BLOCKS.values() for line in block}
    for p in list(doc.paragraphs):
        if p.text.strip() in texts_to_remove:
            remove_paragraph(p)


def insert_section(doc: Document, anchor_body_text: str, heading_text: str):
    anchor = find_paragraph(doc, anchor_body_text)
    heading_style = find_paragraph(doc, heading_text.rsplit(" ", 1)[0] if False else heading_text).style if any(
        p.text.strip() == heading_text for p in doc.paragraphs
    ) else None
    if heading_style is None:
        # fallback: use next heading of same chapter style when exact heading doesn't exist yet
        heading_style = find_paragraph(doc, "3.4.3 任务上传流程").style
        if heading_text.startswith("4."):
            heading_style = find_paragraph(doc, "4.5.5 多点任务场景").style
        if heading_text.startswith("5."):
            heading_style = find_paragraph(doc, "5.4.4 场景分析中的 ROS 含义").style
    body_style = anchor.style

    current = anchor
    block = SECTION_BLOCKS[heading_text]
    for idx, text in enumerate(block):
        style = heading_style if idx == 0 else body_style
        current = insert_after(current, text, style)


def main():
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(str(DOC_PATH))
    cleanup_existing_sections(doc)

    insert_section(
        doc,
        "任务上传是系统由规划阶段转入执行阶段的关键环节。系统在生成各机任务点序列后，通过 MAVROS 服务接口逐机清空旧任务并上传新任务，随后统计各机上传结果，形成任务上传成功率信息。该流程既是系统正常运行的重要基础，也为任务可执行性评价提供了直接的数据来源。",
        "3.4.4 不同高度任务分层规划策略",
    )
    insert_section(
        doc,
        "多点任务场景用于模拟一次飞行过程中依次访问多个任务点的情形。该场景强调全局任务组织而不仅仅是单点绕障，适合分析航点连接关系、路径折返特征和综合指标变化。相较于城市障碍场和通道场，多点任务场更能体现系统在全局任务结构组织方面的能力。",
        "4.5.6 不同高度任务场景",
    )
    insert_section(
        doc,
        "编队保持能力通过对编队误差进行归一化处理形成直观评价值，用于辅助描述编队保持效果。虽然该指标仍具有较强工程性，但对于实验展示具有一定辅助意义。",
        "4.7.9 高度分层规划相关指标",
    )
    insert_section(
        doc,
        "，典型场景结果不仅是路径规划结果的差异，也体现了系统内部节点协同方式的差异。城市障碍场与通道场更多考验 `avoidance_2d` 对安全锚点的修正与 `formation_controller` 对目标队形的重计算能力；多点任务场更考验 `mission_ui` 对多航点任务组织和 MAVROS 航点上传的稳定性；动态改编队场则同时涉及话题广播、任务刷新和可视化重构。换言之，场景越复杂，系统结果越能体现 ROS 节点化组织的必要性。",
        "5.4.5 不同高度分层规划的扩展分析",
    )

    doc.save(str(DOC_PATH))


if __name__ == "__main__":
    main()
