from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT_DESKTOP = Path.home() / "Desktop"
LUNWEN_DIR = ROOT_DESKTOP / "lunwen"
OUTPUT_PATH = LUNWEN_DIR / "毕业论文正式稿_模板整理版.docx"
SOURCE_MD_GLOB = "*第三阶段.md"
TEMPLATE_GLOB = "*模版_202312.docx"


CN_NUM = {
    "零": 0,
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


SKIP_SECTIONS = {"写作说明", "目录", "当前仍需补充的材料清单"}


@dataclass
class ParsedDoc:
    title: str
    abstract: list[str]
    keywords: str
    body_lines: list[str]


def iter_candidates(base: Path, pattern: str) -> Iterable[Path]:
    return sorted(base.rglob(pattern))


def find_first(base: Path, pattern: str) -> Path:
    matches = list(iter_candidates(base, pattern))
    if not matches:
        raise FileNotFoundError(f"未找到匹配文件: {pattern}")
    return matches[0]


def read_markdown(md_path: Path) -> ParsedDoc:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    title = "基于数字孪生精准建模的无人机航路动态规划研究"
    abstract_lines: list[str] = []
    keywords = ""
    body_lines: list[str] = []

    current_heading = None
    skip_mode = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("## "):
            current_heading = stripped[3:].strip()
            skip_mode = current_heading in SKIP_SECTIONS

        if stripped == "## 题目":
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                title = lines[j].strip()
            i = j
            continue

        if stripped == "## 摘要":
            i += 1
            while i < len(lines) and not lines[i].startswith("## "):
                if lines[i].strip():
                    abstract_lines.append(lines[i].strip())
                i += 1
            continue

        if stripped == "## 关键词":
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                keywords = lines[j].strip()
            i = j + 1
            continue

        if stripped.startswith("# ") and "毕业论文正文版" in stripped:
            i += 1
            continue

        if skip_mode:
            i += 1
            continue

        if stripped == "## 题目" or stripped == "## 摘要" or stripped == "## 关键词":
            i += 1
            continue

        if stripped == "## 当前仍需补充的材料清单":
            break

        body_lines.append(line)
        i += 1

    return ParsedDoc(
        title=title,
        abstract=abstract_lines,
        keywords=keywords,
        body_lines=body_lines,
    )


def set_run_font(run, size_pt: float = 12, bold: bool | None = None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, size_pt: float = 12, bold: bool = False):
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold)
    return run


def configure_document(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.8)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name, font_size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        style.font.size = Pt(font_size)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size_pt=10.5)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "打开文档后请在目录处右键选择“更新域”。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size_pt=12)


def chinese_num_to_int(text: str) -> int:
    if text == "十":
        return 10
    if "十" in text:
        parts = text.split("十")
        tens = CN_NUM.get(parts[0], 1) if parts[0] else 1
        ones = CN_NUM.get(parts[1], 0) if len(parts) > 1 else 0
        return tens * 10 + ones
    return CN_NUM.get(text, 0)


def heading_level(line: str) -> tuple[int, str] | None:
    if line.startswith("### "):
        return 3, line[4:].strip()
    if line.startswith("## "):
        return 2, line[3:].strip()
    if line.startswith("# "):
        return 1, line[2:].strip()
    return None


def caption_for_image(alt_text: str, chapter_no: int, figure_index: int) -> str:
    clean = alt_text.strip() if alt_text.strip() else f"图像{figure_index}"
    return f"图{chapter_no}-{figure_index} {clean}"


def add_cover(doc: Document, thesis_title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.space_before = Pt(90)
    set_paragraph_text(p, "重庆大学普通本科毕业论文（设计）", size_pt=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    set_paragraph_text(p, thesis_title, size_pt=20, bold=True)

    cover_lines = [
        "学生姓名：陈银",
        "学    号：需补充",
        "学    院：需补充",
        "专    业：需补充",
        "指导教师：需补充",
        "完成日期：2026年4月",
    ]
    for line in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        set_paragraph_text(p, line, size_pt=14)

    doc.add_page_break()


def add_abstract(doc: Document, abstract_lines: list[str], keywords: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(p, "摘  要", size_pt=16, bold=True)

    for line in abstract_lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.85)
        p.paragraph_format.line_spacing = 1.5
        set_paragraph_text(p, line)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run("关键词：")
    set_run_font(run, size_pt=12, bold=True)
    run = p.add_run(keywords)
    set_run_font(run, size_pt=12)

    doc.add_page_break()


def add_toc_page(doc: Document):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(p, "目  录", size_pt=16, bold=True)
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.5
    add_toc(toc_p)
    doc.add_page_break()


def finalize_footer(doc: Document):
    for section in doc.sections:
        footer_p = section.footer.paragraphs[0]
        add_page_number(footer_p)


def render_body(doc: Document, lines: list[str]):
    current_chapter = 0
    figure_counter = 0
    buffer: list[str] = []

    def flush_buffer():
        nonlocal buffer
        if not buffer:
            return
        text = " ".join(item.strip() for item in buffer if item.strip())
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.85)
            p.paragraph_format.line_spacing = 1.5
            set_paragraph_text(p, text)
        buffer = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_buffer()
            alt_text, img_path = image_match.groups()
            img_file = Path(img_path)
            if img_file.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_file), width=Cm(14.5))
                figure_counter += 1
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(6)
                caption = caption_for_image(alt_text, current_chapter or 1, figure_counter)
                set_paragraph_text(cp, caption, size_pt=10.5)
            continue

        if not stripped:
            flush_buffer()
            continue

        heading = heading_level(stripped)
        if heading:
            flush_buffer()
            level, text = heading
            if level == 1:
                if doc.paragraphs:
                    doc.add_page_break()
                m = re.search(r"第([一二三四五六七八九十]+)章", text)
                if m:
                    current_chapter = chinese_num_to_int(m.group(1))
                else:
                    current_chapter += 1
                figure_counter = 0
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {level}"]
            set_paragraph_text(
                p,
                text,
                size_pt={1: 16, 2: 14, 3: 12}[level],
                bold=True,
            )
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_buffer()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.5
            set_paragraph_text(p, re.sub(r"^\d+\.\s+", "", stripped))
            continue

        if stripped.startswith("- "):
            flush_buffer()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            set_paragraph_text(p, stripped[2:])
            continue

        buffer.append(stripped)

    flush_buffer()


def build_docx():
    md_path = find_first(LUNWEN_DIR, SOURCE_MD_GLOB)
    parsed = read_markdown(md_path)

    doc = Document()
    configure_document(doc)
    add_cover(doc, parsed.title)
    add_abstract(doc, parsed.abstract, parsed.keywords)
    add_toc_page(doc)
    render_body(doc, parsed.body_lines)
    finalize_footer(doc)

    LUNWEN_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_PATH
    try:
        doc.save(output_path)
    except PermissionError:
        output_path = LUNWEN_DIR / "毕业论文正式稿_模板整理版_更新版.docx"
        doc.save(output_path)

    template_matches = list(iter_candidates(ROOT_DESKTOP, TEMPLATE_GLOB))
    if template_matches:
        print(f"Template candidate found: {template_matches[0]}")
    print(f"Markdown source: {md_path}")
    print(f"Output docx: {output_path}")


if __name__ == "__main__":
    build_docx()
